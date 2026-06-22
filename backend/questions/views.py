from django.conf import settings
from django.shortcuts import get_object_or_404
from rest_framework import status as http_status
from rest_framework.decorators import api_view, parser_classes, permission_classes, throttle_classes
from rest_framework.parsers import FormParser, JSONParser, MultiPartParser
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework.throttling import UserRateThrottle

from accounts.models import AuditLog
from centers.models import CenterMembership

from .ai_generation import generate_questions, explain_question_ai, review_code_submission
from .models import Question
from .serializers import QuestionSerializer


class AiQuestionRateThrottle(UserRateThrottle):
    scope = 'ai_question'


class CodeReviewRateThrottle(UserRateThrottle):
    """IT kod savolini AI baholash uchun 'code_review' scope (10/hour)."""
    scope = 'code_review'


class CodeRunRateThrottle(UserRateThrottle):
    """Judge0 kod runner ("Ishga tushirish") uchun 'code_run' scope (20/hour)."""
    scope = 'code_run'


class AiExplainRateThrottle(UserRateThrottle):
    """AI tushuntirish endpoint'i uchun 'ai' scope (settings'da 10/day).

    Eslatma: @api_view function-based view'da `throttle_scope` atributi
    WrappedAPIView'ga ko'chmaganligi va ScopedRateThrottle.allow_request
    scope'ni view'dan qayta o'qiganligi sababli, scope'ni shu yerda
    UserRateThrottle subclass orqali (per-foydalanuvchi) o'rnatamiz —
    AiQuestionRateThrottle bilan bir xil naqsh.
    """
    scope = 'ai'


def _user_can_create_for_center(user, center_id):
    """Teacher/Manager/Owner with approved membership can create questions."""
    if user.is_platform_admin:
        return True
    return CenterMembership.objects.filter(
        user=user, center_id=center_id,
        role__in=[
            CenterMembership.ROLE_TEACHER,
            CenterMembership.ROLE_MANAGER,
            CenterMembership.ROLE_OWNER,
        ],
        status=CenterMembership.STATUS_APPROVED,
    ).exists()


@api_view(['GET', 'POST'])
@permission_classes([IsAuthenticated])
@parser_classes([JSONParser, MultiPartParser, FormParser])
def questions_list_create(request):
    """GET /api/questions/?center=<id>  — list questions for a center.
    POST /api/questions/                 — create one (approved teacher/manager/owner only).
    """
    if request.method == 'GET':
        raw_center = request.query_params.get('center')
        if not raw_center:
            return Response(
                {'detail': 'center query parametri majburiy'},
                status=http_status.HTTP_400_BAD_REQUEST,
            )
        # `?center=abc` kabi noto'g'ri qiymat berilsa, avval bu joyda
        # DB darajasida ValueError otilardi va 500 qaytarardi. Endi 400.
        try:
            center_id = int(raw_center)
        except (TypeError, ValueError):
            return Response(
                {'detail': "center parametri son bo'lishi kerak"},
                status=http_status.HTTP_400_BAD_REQUEST,
            )
        if not _user_can_create_for_center(request.user, center_id):
            return Response(
                {'detail': 'Forbidden'},
                status=http_status.HTTP_403_FORBIDDEN,
            )
        qs = Question.objects.filter(center_id=center_id)
        # Pagination: bitta markazda yuzlab savol to'planishi mumkin — butun
        # ro'yxatni bitta response'da uzatish xotira/trafik jihatdan og'ir.
        # olympiads_list_create kabi LargePageNumberPagination ishlatamiz:
        # frontend `?page_size=200` yuboradi (api.js getQuestions), shu sababli
        # bitta sahifada hammasi keladi va round-trip ko'paymaydi.
        from olympy_api.pagination import LargePageNumberPagination
        paginator = LargePageNumberPagination()
        page = paginator.paginate_queryset(qs, request)
        if page is not None:
            return paginator.get_paginated_response(
                QuestionSerializer(page, many=True, context={'request': request}).data
            )
        return Response(QuestionSerializer(qs, many=True, context={'request': request}).data)

    serializer = QuestionSerializer(data=request.data, context={'request': request})
    serializer.is_valid(raise_exception=True)
    center_id = serializer.validated_data['center'].id
    if not _user_can_create_for_center(request.user, center_id):
        return Response(
            {'detail': "Savol yaratish uchun o'qituvchi/manager arizangiz tasdiqlanishi kerak"},
            status=http_status.HTTP_403_FORBIDDEN,
        )
    question = serializer.save(created_by=request.user)
    AuditLog.log(request, 'question_create', target=question, extra={
        'center_id': question.center_id,
        'subject': getattr(question, 'subject', ''),
    })
    return Response(QuestionSerializer(question, context={'request': request}).data,
                    status=http_status.HTTP_201_CREATED)


@api_view(['GET', 'PATCH', 'PUT', 'DELETE'])
@permission_classes([IsAuthenticated])
@parser_classes([JSONParser, MultiPartParser, FormParser])
def question_detail(request, question_id):
    """GET/PATCH/PUT/DELETE /api/questions/{id}/

    Edit (qalam) tugmasi avval ulanmagan edi — endi PATCH bilan ishlaydi.
    DELETE ham qo'shildi: o'sha ruxsatga ega rolllar.
    """
    question = get_object_or_404(Question, pk=question_id)
    if not _user_can_create_for_center(request.user, question.center_id):
        return Response({'detail': 'Forbidden'}, status=http_status.HTTP_403_FORBIDDEN)
    if request.method == 'GET':
        return Response(QuestionSerializer(question, context={'request': request}).data)
    if request.method == 'DELETE':
        # Audit yozuvini o'chirishdan OLDIN yozamiz — delete() dan keyin pk
        # None bo'lib qoladi va target_id yo'qoladi.
        AuditLog.log(request, 'question_delete', target=question, extra={
            'center_id': question.center_id,
            'subject': getattr(question, 'subject', ''),
        })
        question.delete()
        return Response(status=http_status.HTTP_204_NO_CONTENT)
    # PATCH / PUT
    partial = request.method == 'PATCH'
    data = request.data
    # center maydonini tahrirlashga ruxsat bermaymiz — savol bir markazga
    # bog'liq bo'lib qoladi.
    if hasattr(data, 'copy'):
        data = data.copy()
    else:
        data = dict(data)
    data.pop('center', None)
    serializer = QuestionSerializer(question, data=data, partial=partial, context={'request': request})
    serializer.is_valid(raise_exception=True)
    serializer.save()
    return Response(QuestionSerializer(question, context={'request': request}).data)


@api_view(['DELETE'])
@permission_classes([IsAuthenticated])
def delete_all_questions(request):
    """DELETE /api/questions/delete-all/?center=<id>
    Delete all questions for a center (approved teacher/manager/owner only).
    """
    raw_center = request.query_params.get('center')
    if not raw_center:
        return Response(
            {'detail': 'center query parametri majburiy'},
            status=http_status.HTTP_400_BAD_REQUEST,
        )
    try:
        center_id = int(raw_center)
    except (TypeError, ValueError):
        return Response(
            {'detail': "center parametri son bo'lishi kerak"},
            status=http_status.HTTP_400_BAD_REQUEST,
        )
    if not _user_can_create_for_center(request.user, center_id):
        return Response(
            {'detail': 'Forbidden'},
            status=http_status.HTTP_403_FORBIDDEN,
        )
    ids_raw = request.query_params.get('ids')
    selected_ids = None
    if ids_raw:
        try:
            ids = [int(x) for x in ids_raw.split(',') if x.strip()]
        except (ValueError, TypeError):
            return Response(
                {'detail': "ids parametri butun sonlar ro'yxati bo'lishi kerak"},
                status=http_status.HTTP_400_BAD_REQUEST,
            )
        selected_ids = ids
        deleted_count, _ = Question.objects.filter(center_id=center_id, id__in=ids).delete()
    else:
        deleted_count, _ = Question.objects.filter(center_id=center_id).delete()
    # Ommaviy o'chirish ham audit'ga yoziladi (yakka o'chirish question_delete
    # bilan log qilingani kabi). Ko'p obyekt o'chirilgani uchun target=None;
    # detallar extra'da: markaz, o'chirilgan soni va (qisman bo'lsa) ID'lar.
    AuditLog.log(request, 'question_bulk_delete', extra={
        'center_id': center_id,
        'deleted_count': deleted_count,
        'ids': selected_ids,
    })
    return Response(
        {'detail': f"{deleted_count} ta savol muvaffaqiyatli o'chirildi", 'deleted_count': deleted_count},
        status=http_status.HTTP_200_OK,
    )


@api_view(['POST'])
@permission_classes([IsAuthenticated])
@throttle_classes([AiQuestionRateThrottle])
def generate_ai_questions(request):
    """POST /api/questions/generate-ai/ — preview AI questions before saving."""
    center_id = request.data.get('center')
    if not center_id:
        return Response(
            {'detail': 'center majburiy'},
            status=http_status.HTTP_400_BAD_REQUEST,
        )
    if not _user_can_create_for_center(request.user, center_id):
        return Response(
            {'detail': "Savol yaratish uchun o'qituvchi/manager arizangiz tasdiqlanishi kerak"},
            status=http_status.HTTP_403_FORBIDDEN,
        )

    # Premium + oylik limit check. SubscriptionService bir joyda premium
    # holatini va plandagi oylik AI generatsiya limitini (Standart 20, Plus
    # 100, Pro cheksiz) tekshiradi.
    from centers.models import EducationCenter
    from billing.services import SubscriptionService
    center = EducationCenter.objects.filter(pk=center_id).first()
    svc = SubscriptionService(center) if center else None
    if svc and not svc.is_premium:
        return Response(
            {
                'detail': "AI yordamida savol yaratish faqat premium tashkilotlar uchun. Premium obunani faollashtiring.",
                'upgrade_required': True
            },
            status=http_status.HTTP_403_FORBIDDEN
        )
    if svc and not svc.can_use_ai_generation():
        limit = svc.ai_generation_monthly_limit
        return Response(
            {
                'detail': f"Ushbu oy uchun AI savol generatsiya limiti ({limit}) tugadi. Tarifni yangilang yoki keyingi oyni kuting.",
                'upgrade_required': True,
                'limit_reached': True,
            },
            status=http_status.HTTP_403_FORBIDDEN
        )

    result = generate_questions(
        subject=request.data.get('subject'),
        topic=request.data.get('topic'),
        count=request.data.get('count', 10),
        difficulty=request.data.get('difficulty', 'medium'),
        question_type=request.data.get('question_type'),
    )
    if not result.get('ok'):
        status_code = (
            http_status.HTTP_400_BAD_REQUEST
            if result.get('error') in ("Fan va mavzu majburiy.",)
            else http_status.HTTP_503_SERVICE_UNAVAILABLE
        )
        return Response(
            {'detail': result.get('error') or "AI savol yarata olmadi"},
            status=status_code,
        )
    # Muvaffaqiyatli generatsiyani oylik limit hisobiga yozamiz (faqat haqiqiy
    # natija bo'lganda — xato/bo'sh javob limitdan ushlab qolinmaydi).
    if svc:
        try:
            svc.log_ai_generation(user=request.user, count=len(result.get('questions') or []))
        except Exception:
            pass
    return Response({'questions': result['questions']})


def _normalize_correct_answer(value):
    """To'g'ri javob qiymatini 0-based indeksga aylantiradi. Noma'lumda None.

    Qo'llab-quvvatlanadigan ko'rinishlar (katta/kichik harf farqsiz):
      - A/B/C/D/E/F        → 0..5
      - 0..9               → o'sha son
      - "1.", "2." …       → 0,1 … (nuqtali tartib raqami, 1-based deb olinadi)
      - "1-variant", "2 variant", "variant a", "variant 1" → 0,1 …
      - "birinchi"/"ikkinchi"/"uchinchi"/"to'rtinchi"      → 0,1,2,3
    Foydalanuvchilar Word/Excel jadvallarida to'g'ri javobni juda turli
    ko'rinishda yozadilar — shu sababli keng tahlil qilamiz.
    """
    if value is None:
        return None
    s = str(value).strip()
    if not s:
        return None
    up = s.upper()

    letter_map = {'A': 0, 'B': 1, 'C': 2, 'D': 3, 'E': 4, 'F': 5}
    if up in letter_map:
        return letter_map[up]

    # Sof son: "0", "1" … (0-based deb qabul qilamiz — eski xulq saqlanadi).
    try:
        i = int(up)
        if 0 <= i <= 9:
            return i
    except (TypeError, ValueError):
        pass

    low = s.lower()

    # So'z bilan yozilgan tartib ("birinchi" … "to'rtinchi"). Apostrofning
    # turli ko'rinishlarini (' ` ’) bir xil normaga keltiramiz.
    low_norm = low.replace('`', "'").replace('’', "'")
    word_map = {
        'birinchi': 0, 'ikkinchi': 1, 'uchinchi': 2,
        "to'rtinchi": 3, 'beshinchi': 4, 'oltinchi': 5,
    }
    if low_norm in word_map:
        return word_map[low_norm]

    # "variant a" / "variant b" ko'rinishi — oxiridagi harfdan indeks.
    if low_norm.startswith('variant'):
        tail = low_norm[len('variant'):].strip(' .-:').upper()
        if tail in letter_map:
            return letter_map[tail]
        # "variant 1" / "variant 2" — 1-based tartib raqami.
        try:
            n = int(tail)
            if 1 <= n <= 9:
                return n - 1
        except (TypeError, ValueError):
            pass

    # "4-B", "9-A", "4.B", "4/B" — son+ajratuvchi+harf: harfni olamiz.
    import re
    m = re.match(r'^\s*\d+\s*[-./]\s*([A-Fa-f])\s*$', s, re.IGNORECASE)
    if m:
        return letter_map.get(m.group(1).upper())

    # Boshida tartib raqami bo'lgan ko'rinishlar: "1.", "2)", "1-variant",
    # "2 variant", "1-javob" … — birinchi sonni ajratib, 1-based deb olamiz.
    m = re.match(r'^\s*(\d+)', low_norm)
    if m:
        n = int(m.group(1))
        if 1 <= n <= 9:
            return n - 1

    # So'nggi urinish: "(A)", "javob: B", "A)", "[C]", "to'g'ri: A" kabi
    # izolyatsiyalangan A-F harfni topamiz — harf atrofida harf bo'lmasin.
    m = re.search(r'(?:^|[^a-zA-Z])([A-Fa-f])(?:[^a-zA-Z]|$)', s)
    if m:
        return letter_map.get(m.group(1).upper())

    return None


def _normalize_difficulty(value):
    if not value:
        return Question.DIFFICULTY_MEDIUM
    s = str(value).strip().lower()
    aliases = {
        'oson': 'easy', 'easy': 'easy', 'beginner': 'beginner', 'elementary': 'elementary',
        "o'rta": 'medium', "orta": 'medium', "o`rta": 'medium', 'medium': 'medium', 'int': 'int', 'intermediate': 'int',
        'qiyin': 'hard', 'hard': 'hard', 'advanced': 'advanced',
        'pre-int': 'pre-int', 'pre-intermediate': 'pre-int', 'upper-int': 'upper-int', 'upper-intermediate': 'upper-int',
    }
    return aliases.get(s, Question.DIFFICULTY_MEDIUM)


# Excel/CSV va Word import uchun umumiy ustun sarlavhalari va namuna qatori.
# Bitta joyda saqlanadi — Excel namuna (frontend CSV), Word shablon va
# hujjatlardagi format tavsifi bir xil bo'lib qolishi uchun.
IMPORT_HEADER = ['savol', 'variant_a', 'variant_b', 'variant_c', 'variant_d', 'togri_javob', 'qiyinlik', 'fan']
IMPORT_SAMPLE_ROW = ['2+2 nechaga teng?', '3', '4', '5', '6', 'B', 'easy', 'Matematika']


def _row_looks_like_header(row):
    """Birinchi qator sarlavhami? togri_javob ustunida A-F/0-5 emas, matn bo'lsa header.

    5-ustunli format (savol|A|B|C|togri_javob) va 6+ ustunli format ikkalasini
    qo'llab-quvvatlaydi. togri_javob indeksi: 5-col→4, 6+col→5.
    """
    if not row or len(row) < 5:
        return False
    check_idx = 5 if len(row) >= 6 else 4
    s = str(row[check_idx] or '').strip().upper()
    if not s:
        return True
    if s in ('A', 'B', 'C', 'D', 'E', 'F', '0', '1', '2', '3', '4', '5'):
        return False
    return True


# Header orqali ustunlarni aniqlash uchun alias ro'yxati. Foydalanuvchilar
# o'z Word/Excel jadvallarini turli ustun nomlari va tartibi bilan yaratadilar;
# birinchi qator sarlavha bo'lsa, ustun nomini pastdagi aliaslar bilan
# solishtirib mapping quramiz (katta/kichik harfga e'tibor bermay).
COLUMN_ALIASES = {
    'text':       ['savol', 'savol matni', 'savol_matni', 'question', 'matn', 'text'],
    'option_a':   ['variant_a', 'a', 'a_variant', 'javob_a', 'a variant', '1-variant', '1 variant'],
    'option_b':   ['variant_b', 'b', 'b_variant', 'javob_b', 'b variant', '2-variant', '2 variant'],
    'option_c':   ['variant_c', 'c', 'c_variant', 'javob_c', 'c variant', '3-variant', '3 variant'],
    'option_d':   ['variant_d', 'd', 'd_variant', 'javob_d', 'd variant', '4-variant', '4 variant'],
    'correct':    ['togri_javob', "to'g'ri_javob", "to'g'ri javob", 'javob', 'answer', 'correct', 'togri javob', "to'g'ri"],
    'difficulty': ['qiyinlik', 'difficulty', 'daraja', 'level'],
    'subject':    ['fan', 'subject', 'fani', 'mavzu'],
}


def _detect_column_map(header_row):
    """Sarlavha qatoridan {maydon: indeks} mapping quradi.

    Har bir katakni COLUMN_ALIASES bilan solishtiradi (strip + lower). Topilgan
    maydonlarni qaytaradi. Minimal talab (text, correct, kamida 2 ta option)
    bajarilmasa — None qaytaradi, shunda chaqiruvchi positional mapping'ga
    qaytadi. Bir aliasga bir nechta katak mos kelsa, birinchisi olinadi.
    """
    if not header_row:
        return None
    cells = [('' if c is None else str(c).strip().lower()) for c in header_row]
    mapping = {}
    for field, aliases in COLUMN_ALIASES.items():
        alias_set = {a.strip().lower() for a in aliases}
        for i, cell in enumerate(cells):
            if cell and cell in alias_set and field not in mapping:
                mapping[field] = i
                break
    option_count = sum(1 for k in ('option_a', 'option_b', 'option_c', 'option_d') if k in mapping)
    if 'text' in mapping and 'correct' in mapping and option_count >= 2:
        return mapping
    return None


# Data-based aniqlash uchun qiyinlik so'zlari (header bo'lmaganda ustun rolini
# topish uchun). _normalize_difficulty aliaslari bilan mos.
_DIFFICULTY_TOKENS = {
    'easy', 'medium', 'hard', 'beginner', 'elementary', 'intermediate',
    'advanced', 'int', 'pre-int', 'upper-int',
    'oson', "o'rta", 'orta', "o`rta", 'qiyin',
}


def _infer_columns_from_data(rows):
    """Ustun nomi (header) bo'lmaganda, ma'lumotning o'zidan ustun rolini topadi.

    Har bir ustunni barcha qatorlar bo'yicha skanlaymiz va shu mezonlar bilan
    rol beramiz:
      - correct:    qiymatlar deyarli butunlay A/B/C/D yoki 0/1/2/3 (ya'ni
                    _normalize_correct_answer tanийdigan) bo'lsa
      - difficulty: qiymatlar deyarli butunlay qiyinlik so'zlari bo'lsa
      - subject:    qisqa (1-3 so'z), takrorlanuvchi qiymatli ustun (fan nomi
                    odatda butun faylda bir xil yoki bir nechta marta takror)
      - text:       o'rtacha matn uzunligi eng katta ustun (savol matni)
      - options:    qolgan ustunlar, chapdan o'ngga A,B,C,D tartibida

    Minimal talab: text + correct + kamida 2 ta option topilsa — col_map dict
    qaytaradi (xuddi _detect_column_map kabi {field: index}). Aks holda None.
    Header tashlanmaydi — bu rejim faqat header yo'qligida chaqiriladi, demak
    barcha qatorlar ma'lumot.
    """
    if not rows:
        return None

    # Ustunlar sonini eng keng qatorga qarab aniqlaymiz.
    n_cols = max((len(r) for r in rows if r), default=0)
    if n_cols < 4:
        # text + correct + 2 option uchun kamida 4 ustun kerak.
        return None

    # Har bir ustun bo'yicha statistika yig'amiz.
    col_values = [[] for _ in range(n_cols)]
    for row in rows:
        if not row:
            continue
        for c in range(n_cols):
            val = '' if c >= len(row) or row[c] is None else str(row[c]).strip()
            col_values[c].append(val)

    col_stats = []
    for c in range(n_cols):
        vals = [v for v in col_values[c] if v]
        total = len(vals)
        if total == 0:
            col_stats.append({
                'index': c, 'total': 0, 'correct_ratio': 0.0,
                'letter_ratio': 0.0, 'difficulty_ratio': 0.0, 'avg_len': 0.0,
                'distinct_ratio': 1.0, 'short_ratio': 0.0, 'question_ratio': 0.0,
            })
            continue
        correct_hits = sum(1 for v in vals if _normalize_correct_answer(v) is not None)
        # Harfiy (A-F) correct belgisi: sof sonli option ustunidan (variantlar
        # 1,2,3,4 bo'lishi mumkin) farqlash uchun alohida hisoblaymiz.
        letter_hits = sum(1 for v in vals if v.strip().upper() in ('A', 'B', 'C', 'D', 'E', 'F'))
        diff_hits = sum(1 for v in vals if v.lower().replace('`', "'").replace('’', "'") in _DIFFICULTY_TOKENS)
        avg_len = sum(len(v) for v in vals) / total
        distinct_ratio = len(set(v.lower() for v in vals)) / total
        short_hits = sum(1 for v in vals if len(v.split()) <= 3 and len(v) <= 30)
        # "Savolga o'xshashlik": savol matni odatda jumla — ko'p so'zli yoki '?'
        # bilan tugaydi. Fan nomi esa bitta so'z. Bu signal fan ustunidan qisqaroq
        # savolni (qisqa fayllarda) ham to'g'ri ajratishga yordam beradi.
        question_hits = sum(1 for v in vals if ('?' in v) or len(v.split()) >= 3)
        col_stats.append({
            'index': c,
            'total': total,
            'correct_ratio': correct_hits / total,
            'letter_ratio': letter_hits / total,
            'difficulty_ratio': diff_hits / total,
            'avg_len': avg_len,
            'distinct_ratio': distinct_ratio,
            'short_ratio': short_hits / total,
            'question_ratio': question_hits / total,
        })

    assigned = {}  # index -> field, qaysi ustun band bo'lganini kuzatamiz
    mapping = {}

    # 1) correct ustun. Eng ishonchli belgi — harfiy A/B/C/D ustun: avval shularni
    #    ko'rib chiqamiz. Sof sonli ustun ham correct bo'lishi mumkin (0/1/2/3),
    #    lekin variantlar ham sof son bo'lishi mumkin ('3','4','5'), shuning uchun
    #    sonli holatda ustunni eng o'ng tomondan tanlaymiz — to'g'ri javob odatda
    #    variantlardan keyin, oxirgi ustunlarda turadi.
    letter_candidates = [s for s in col_stats if s['letter_ratio'] >= 0.8]
    if letter_candidates:
        # Bir nechta harfiy ustun bo'lsa (kam uchraydi), eng o'ngdagisini olamiz.
        letter_candidates.sort(key=lambda s: (-s['letter_ratio'], -s['index']))
        ci = letter_candidates[0]['index']
        mapping['correct'] = ci
        assigned[ci] = 'correct'
    else:
        correct_candidates = [s for s in col_stats if s['correct_ratio'] >= 0.8]
        if correct_candidates:
            # Sonli variantlar (1,2,3,4) ham correct_ratio yuqori chiqishi mumkin.
            # Haqiqiy correct ustun: eng qisqa qiymatli (avg_len kichik) va o'ng tomonda.
            correct_candidates.sort(key=lambda s: (-s['correct_ratio'], s['avg_len'], -s['index']))
            ci = correct_candidates[0]['index']
            mapping['correct'] = ci
            assigned[ci] = 'correct'

    # 2) difficulty ustun: difficulty_ratio >= 0.6 bo'lgan, hali band bo'lmagan.
    diff_candidates = [s for s in col_stats if s['index'] not in assigned and s['difficulty_ratio'] >= 0.6]
    if diff_candidates:
        diff_candidates.sort(key=lambda s: -s['difficulty_ratio'])
        di = diff_candidates[0]['index']
        mapping['difficulty'] = di
        assigned[di] = 'difficulty'

    # 3) subject ustun: text'dan OLDIN aniqlaymiz — aks holda takrorlanuvchi fan
    #    nomi (savol matnidan uzunroq bo'lsa) noto'g'ri text deb tanlanishi mumkin.
    #    Fan nomi butun faylda bir xil yoki kam o'zgaradi. Faqat takror bo'lsa
    #    (distinct_ratio < 0.6) va asosan qisqa (short_ratio >= 0.6) bo'lsa subject.
    #    Savol matnlari odatda har xil (yuqori distinct) — bu mezonга tushmaydi.
    subj_candidates = [
        s for s in col_stats
        if s['index'] not in assigned and s['total'] > 0
        and s['distinct_ratio'] < 0.6 and s['short_ratio'] >= 0.6
    ]
    if subj_candidates:
        # Eng takrorlanuvchi (eng kam distinct) ustun fan bo'lishi ehtimoli yuqori.
        subj_candidates.sort(key=lambda s: (s['distinct_ratio'], s['avg_len']))
        si = subj_candidates[0]['index']
        mapping['subject'] = si
        assigned[si] = 'subject'

    # 4) text ustun: savol matni. Ikki belgi muhim — (a) o'rtacha uzunligi katta,
    #    (b) qiymatlar deyarli har xil (savollar takrorlanmaydi → yuqori
    #    distinct_ratio). Avval distinct (>= 0.5) ustunlar ichidan eng uzunini
    #    tanlaymiz; bunday ustun bo'lmasa — oddiy eng uzun avg_len. Bu fan nomi
    #    (qisqa fayllarda savoldan uzunroq bo'lib qolishi mumkin) bilan savol
    #    matnini chalkashtirib yubormaslik uchun.
    text_pool = [s for s in col_stats if s['index'] not in assigned and s['total'] > 0]
    if text_pool:
        # Avvalo "jumla"ga o'xshagan (question_ratio >= 0.5) distinct ustunlarni
        # ko'rib chiqamiz — savol matni '?' yoki ko'p so'zli bo'ladi, fan emas.
        question_pool = [s for s in text_pool if s['question_ratio'] >= 0.5 and s['distinct_ratio'] >= 0.5]
        if question_pool:
            question_pool.sort(key=lambda s: (-s['question_ratio'], -s['avg_len']))
            ti = question_pool[0]['index']
        else:
            # Jumla signali yo'q — distinct (savollar har xil) ustunlar ichidan
            # eng uzunini, ular ham bo'lmasa oddiy eng uzun avg_len.
            distinct_pool = [s for s in text_pool if s['distinct_ratio'] >= 0.5]
            chosen_pool = distinct_pool or text_pool
            chosen_pool.sort(key=lambda s: -s['avg_len'])
            ti = chosen_pool[0]['index']
        mapping['text'] = ti
        assigned[ti] = 'text'

    # 5) options: qolgan barcha ustunlar, chapdan o'ngga A,B,C,D.
    option_fields = ['option_a', 'option_b', 'option_c', 'option_d']
    remaining = sorted(s['index'] for s in col_stats if s['index'] not in assigned and s['total'] > 0)
    for i, col_idx in enumerate(remaining):
        if i >= len(option_fields):
            break
        mapping[option_fields[i]] = col_idx
        assigned[col_idx] = option_fields[i]

    option_count = sum(1 for k in option_fields if k in mapping)
    if 'text' in mapping and 'correct' in mapping and option_count >= 2:
        return mapping
    return None


def _create_questions_from_rows(rows, center_id, user, fallback_subject):
    """Xom qatorlar ro'yxatidan (har biri ustunlar list'i) Question yaratadi.

    Excel/CSV va Word import bir xil validatsiya va xato handling ishlatishi
    uchun ajratilgan umumiy yadro. Avval bu mantiq faqat import_questions_excel
    ichida edi. Qaytaradi: (created_count, errors_list).
    `rows` — birinchi qator sarlavha bo'lishi mumkin (heuristic'da tashlanadi).

    Ustun aniqlash uch bosqichda:
      1) Header-based: birinchi qator sarlavha bo'lib, nomlar COLUMN_ALIASES
         orqali tanib olinsa — header mapping (ustun tartibi ixtiyoriy).
      2) Data-based: header topilmasa, _infer_columns_from_data barcha
         qatorlardan ustun rolini chiqaradi (ustun nomi shart emas).
      3) Positional: 1 va 2 ishlamasa — qat'iy tartib (6+ ustun talab).
    """
    errors = []
    if not rows:
        return 0, errors

    # Bosqich 1 — header: avval _detect_column_map bilan birinchi qatorni
    # sarlavha sifatida sinab ko'ramiz (ustun soni 6 dan kam bo'lsa ham).
    col_map = _detect_column_map(rows[0])
    if col_map is not None:
        # Header topildi — birinchi qator (sarlavha) tashlanadi.
        data_rows = rows[1:]
        is_header = True
    else:
        # Bosqich 2 — data-based: header topilmadi, ma'lumotning o'zidan ustun
        # rolini chiqarishga harakat qilamiz. MUHIM: bu yerda butun `rows` ustida
        # ishlaymiz va birinchi qatorni TASHLAMAYMIZ — chunki header yo'q, birinchi
        # qator ham haqiqiy savol. (_row_looks_like_header birinchi data qatorini
        # noto'g'ri sarlavha deb tashlab yuborishi mumkin edi.)
        inferred = _infer_columns_from_data(rows)
        if inferred is not None:
            col_map = inferred
            data_rows = rows
            is_header = False
        else:
            # Bosqich 3 — positional fallback: 1 va 2 ishlamadi. Faqat shu yerda
            # eski heuristic bilan birinchi qatorni sarlavha bo'lsa tashlaymiz.
            is_header = _row_looks_like_header(rows[0])
            data_rows = rows[1:] if is_header else rows

    created = 0
    for idx, raw_row in enumerate(data_rows, start=2 if is_header else 1):
        if not raw_row:
            continue
        # Bo'sh qatorlarni o'tkazib yuboramiz.
        normalized = [('' if v is None else str(v).strip()) for v in raw_row]
        if not any(normalized):
            continue

        if col_map is not None:
            # Header mapping rejimi: indekslarni nomlar orqali olamiz.
            def _at(field):
                i = col_map.get(field)
                if i is None or i >= len(normalized):
                    return ''
                return normalized[i]

            text = _at('text')
            if not text:
                errors.append({'row': idx, 'detail': "Savol matni bo'sh"})
                continue
            options_raw = [_at('option_a'), _at('option_b'), _at('option_c'), _at('option_d')]
            options = [o for o in options_raw if o]
            if len(options) < 2:
                errors.append({'row': idx, 'detail': "Kamida 2 ta javob varianti kerak"})
                continue
            correct_raw = _at('correct')
            correct_idx = _normalize_correct_answer(correct_raw)
            if correct_idx is None or correct_idx >= len(options):
                errors.append({'row': idx, 'detail': f"To'g'ri javob ko'rsatkichi noto'g'ri: {correct_raw}"})
                continue
            difficulty = _normalize_difficulty(_at('difficulty'))
            subject = _at('subject').strip() or fallback_subject
        else:
            # Positional rejim: qat'iy tartib bo'yicha mapping.
            # 5-col format: savol | A | B | C | togri_javob  (variant D yo'q)
            # 6+ col format: savol | A | B | C | D | togri_javob | (qiyinlik) | (fan)
            if len(normalized) < 5:
                errors.append({'row': idx, 'detail': (
                    f"Yetarli ustun yo'q: {len(normalized)} ta topildi, kamida 5 ta kerak. "
                    "Format: savol | variant_a | variant_b | variant_c | togri_javob "
                    "(yoki 6-ustunli: savol | A | B | C | D | togri_javob). "
                    "Word namuna shablonini yuklab oling."
                )})
                continue
            text = normalized[0]
            if not text:
                errors.append({'row': idx, 'detail': "Savol matni bo'sh"})
                continue
            if len(normalized) == 5:
                # savol | A | B | C | togri_javob
                options_raw = [normalized[1], normalized[2], normalized[3], '']
                correct_col = 4
            else:
                # savol | A | B | C | D | togri_javob | (qiyinlik) | (fan)
                options_raw = [normalized[1], normalized[2], normalized[3], normalized[4]]
                correct_col = 5
            options = [o for o in options_raw if o]
            if len(options) < 2:
                errors.append({'row': idx, 'detail': "Kamida 2 ta javob varianti kerak"})
                continue
            correct_idx = _normalize_correct_answer(normalized[correct_col] if len(normalized) > correct_col else '')
            if correct_idx is None or correct_idx >= len(options):
                errors.append({'row': idx, 'detail': f"To'g'ri javob ko'rsatkichi noto'g'ri: {normalized[correct_col] if len(normalized) > correct_col else ''}"})
                continue
            difficulty = _normalize_difficulty(normalized[correct_col + 1] if len(normalized) > correct_col + 1 else '')
            subject = (normalized[correct_col + 2] if len(normalized) > correct_col + 2 else '').strip() or fallback_subject

        try:
            Question.objects.create(
                center_id=center_id,
                subject=subject[:80],
                text=text,
                options=options,
                correct_answer=correct_idx,
                score=3,
                difficulty=difficulty,
                source=Question.SOURCE_IMPORT,
                created_by=user,
            )
            created += 1
        except Exception as exc:
            errors.append({'row': idx, 'detail': f"DB xatosi: {exc}"})

    return created, errors


@api_view(['POST'])
@permission_classes([IsAuthenticated])
@parser_classes([MultiPartParser, FormParser])
@throttle_classes([AiQuestionRateThrottle])
def import_questions_excel(request):
    """POST /api/questions/import/?center=<id>

    Excel (.xlsx) yoki CSV (.csv) faylidan savollar import qiladi.
    Format (birinchi qator — sarlavha, e'tiborga olinmaydi):
        savol | variant_a | variant_b | variant_c | variant_d | togri_javob | qiyinlik | fan
    `togri_javob`: A/B/C/D yoki 0/1/2/3.
    `qiyinlik`: easy/medium/hard yoki o'zbek nomlari (Oson/O'rta/Qiyin).
    `fan` bo'sh bo'lsa, ?subject= query parametri yoki "Umumiy" ishlatiladi.
    """
    raw_center = request.query_params.get('center') or request.data.get('center')
    if not raw_center:
        return Response(
            {'detail': 'center query parametri majburiy'},
            status=http_status.HTTP_400_BAD_REQUEST,
        )
    try:
        center_id = int(raw_center)
    except (TypeError, ValueError):
        return Response(
            {'detail': "center parametri son bo'lishi kerak"},
            status=http_status.HTTP_400_BAD_REQUEST,
        )
    if not _user_can_create_for_center(request.user, center_id):
        return Response(
            {'detail': "Savol yaratish uchun o'qituvchi/manager arizangiz tasdiqlanishi kerak"},
            status=http_status.HTTP_403_FORBIDDEN,
        )

    upload = request.FILES.get('file') or request.FILES.get('upload') or request.FILES.get('excel')
    if not upload:
        return Response({'detail': 'Fayl yuboring (form key: file)'}, status=http_status.HTTP_400_BAD_REQUEST)

    # Fayl hajmi cheklovi: avval Excel/CSV import'ga limit yo'q edi va katta
    # fayl (openpyxl butun workbook'ni yuklaydi) xotirani to'ldirishi mumkin edi.
    import_max_bytes = getattr(settings, 'AI_QUESTION_IMPORT_MAX_BYTES', 10 * 1024 * 1024)
    if upload.size and upload.size > import_max_bytes:
        return Response(
            {'detail': f"Fayl juda katta. Limit: {import_max_bytes // (1024 * 1024)} MB"},
            status=http_status.HTTP_400_BAD_REQUEST,
        )

    filename = (getattr(upload, 'name', '') or '').lower()
    fallback_subject = (request.query_params.get('subject') or request.data.get('subject') or 'Umumiy').strip() or 'Umumiy'

    rows = []
    try:
        if filename.endswith('.csv'):
            import csv
            import io
            raw = upload.read()
            # BOM va encoding fallback
            for enc in ('utf-8-sig', 'utf-8', 'cp1251', 'latin-1'):
                try:
                    text = raw.decode(enc)
                    break
                except UnicodeDecodeError:
                    text = None
                    continue
            if text is None:
                return Response({'detail': "CSV fayl encoding'i aniqlanmadi"}, status=http_status.HTTP_400_BAD_REQUEST)
            # Avval ; bilan urinib ko'ramiz, agar bo'sh chiqsa , ga o'tamiz.
            try:
                dialect = csv.Sniffer().sniff(text[:2000], delimiters=',;\t|')
                reader = csv.reader(io.StringIO(text), dialect)
            except Exception:
                reader = csv.reader(io.StringIO(text))
            for r in reader:
                rows.append(r)
        elif filename.endswith('.xlsx') or filename.endswith('.xlsm'):
            try:
                from openpyxl import load_workbook
            except ImportError:
                return Response(
                    {'detail': "openpyxl o'rnatilmagan. Iltimos administratorga xabar bering."},
                    status=http_status.HTTP_503_SERVICE_UNAVAILABLE,
                )
            wb = load_workbook(upload, read_only=True, data_only=True)
            ws = wb.active
            for row in ws.iter_rows(values_only=True):
                rows.append(list(row))
        else:
            return Response(
                {'detail': "Faqat .xlsx yoki .csv fayl qabul qilinadi"},
                status=http_status.HTTP_400_BAD_REQUEST,
            )
    except Exception as exc:
        return Response(
            {'detail': f"Faylni o'qib bo'lmadi: {exc}"},
            status=http_status.HTTP_400_BAD_REQUEST,
        )

    if not rows:
        return Response({'detail': "Faylda satr topilmadi"}, status=http_status.HTTP_400_BAD_REQUEST)

    # Qatorlarni Question'larga aylantirish — Word import bilan umumiy yadro.
    created, errors = _create_questions_from_rows(rows, center_id, request.user, fallback_subject)

    return Response({
        'created': created,
        'errors': errors[:50],  # Frontend'da ko'p xato ko'rsatmaslik uchun cheklov
        'error_count': len(errors),
    })


@api_view(['POST'])
@permission_classes([IsAuthenticated])
@parser_classes([MultiPartParser, FormParser])
@throttle_classes([AiQuestionRateThrottle])
def import_questions_word(request):
    """POST /api/questions/import-word/?center=<id>

    Word (.docx) faylidan savollar import qiladi. Fayl ichida JADVAL (table)
    bo'lishi shart — xuddi Excel shablonidagidek ustunlar bilan:
        savol | variant_a | variant_b | variant_c | variant_d | togri_javob | qiyinlik | fan
    Birinchi qator sarlavha (header) deb tashlanadi, qolgan qatorlar — savollar.
    Erkin matn parse qilinmaydi — faqat jadval o'qiladi.
    Validatsiya va xato handling Excel import bilan bir xil (umumiy yadro).
    """
    raw_center = request.query_params.get('center') or request.data.get('center')
    if not raw_center:
        return Response(
            {'detail': 'center query parametri majburiy'},
            status=http_status.HTTP_400_BAD_REQUEST,
        )
    try:
        center_id = int(raw_center)
    except (TypeError, ValueError):
        return Response(
            {'detail': "center parametri son bo'lishi kerak"},
            status=http_status.HTTP_400_BAD_REQUEST,
        )
    if not _user_can_create_for_center(request.user, center_id):
        return Response(
            {'detail': "Savol yaratish uchun o'qituvchi/manager arizangiz tasdiqlanishi kerak"},
            status=http_status.HTTP_403_FORBIDDEN,
        )

    upload = request.FILES.get('file') or request.FILES.get('upload') or request.FILES.get('word')
    if not upload:
        return Response({'detail': 'Fayl yuboring (form key: file)'}, status=http_status.HTTP_400_BAD_REQUEST)

    # Fayl hajmi cheklovi — Excel import bilan bir xil limit (python-docx butun
    # hujjatni xotiraga yuklaydi).
    import_max_bytes = getattr(settings, 'AI_QUESTION_IMPORT_MAX_BYTES', 10 * 1024 * 1024)
    if upload.size and upload.size > import_max_bytes:
        return Response(
            {'detail': f"Fayl juda katta. Limit: {import_max_bytes // (1024 * 1024)} MB"},
            status=http_status.HTTP_400_BAD_REQUEST,
        )

    filename = (getattr(upload, 'name', '') or '').lower()
    if not filename.endswith('.docx'):
        return Response(
            {'detail': "Faqat .docx fayl qabul qilinadi"},
            status=http_status.HTTP_400_BAD_REQUEST,
        )

    fallback_subject = (request.query_params.get('subject') or request.data.get('subject') or 'Umumiy').strip() or 'Umumiy'

    try:
        from docx import Document
    except ImportError:
        return Response(
            {'detail': "python-docx o'rnatilmagan. Iltimos administratorga xabar bering."},
            status=http_status.HTTP_503_SERVICE_UNAVAILABLE,
        )

    rows = []
    try:
        document = Document(upload)
        # Faqat jadval(lar)dan qatorlarni yig'amiz. Bir nechta jadval bo'lsa
        # hammasini ketma-ket qo'shamiz — birinchi jadvalning header'i qolganlar
        # uchun ham amal qiladi (umumiy yadro birinchi qatorni tashlaydi).
        for table in document.tables:
            for cell_row in table.rows:
                seen_tc = set()
                row_cells = []
                for cell in cell_row.cells:
                    tc_id = id(cell._tc)
                    if tc_id not in seen_tc:
                        seen_tc.add(tc_id)
                        row_cells.append(cell.text.strip())
                if row_cells:
                    rows.append(row_cells)
    except Exception as exc:
        return Response(
            {'detail': f"Word faylni o'qib bo'lmadi: {exc}"},
            status=http_status.HTTP_400_BAD_REQUEST,
        )

    if not rows:
        return Response(
            {'detail': "Word faylda jadval topilmadi. Savollar jadval (table) ko'rinishida bo'lishi kerak — namunani yuklab oling."},
            status=http_status.HTTP_400_BAD_REQUEST,
        )

    # Qatorlarni Question'larga aylantirish — Excel import bilan umumiy yadro.
    created, errors = _create_questions_from_rows(rows, center_id, request.user, fallback_subject)

    return Response({
        'created': created,
        'errors': errors[:50],  # Frontend'da ko'p xato ko'rsatmaslik uchun cheklov
        'error_count': len(errors),
    })


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def download_word_template(request):
    """GET /api/questions/word-template/ — savol import uchun namuna .docx fayl.

    Foydalanuvchi yuklab oladi, jadvalni to'ldiradi va import-word/ ga yuklaydi.
    Excel namunasi (frontend CSV) bilan bir xil ustunlar — IMPORT_HEADER.
    """
    try:
        from docx import Document
    except ImportError:
        return Response(
            {'detail': "python-docx o'rnatilmagan. Iltimos administratorga xabar bering."},
            status=http_status.HTTP_503_SERVICE_UNAVAILABLE,
        )

    import io

    from django.http import HttpResponse

    document = Document()
    document.add_heading('Olympy — savollar import shabloni', level=1)
    document.add_paragraph(
        "Quyidagi jadvalni to'ldiring. Birinchi qator (sarlavha) o'zgartirilmasin. "
        "Har bir qatorga bitta savol yozing. To'g'ri javob: A/B/C/D yoki 0/1/2/3. "
        "Qiyinlik: easy/medium/hard yoki Oson/O'rta/Qiyin. "
        "Fan bo'sh bo'lsa, import paytidagi fan yoki \"Umumiy\" ishlatiladi."
    )

    table = document.add_table(rows=1, cols=len(IMPORT_HEADER))
    try:
        table.style = 'Table Grid'
    except Exception:
        # Ba'zi shablon-siz hujjatlarda 'Table Grid' stili bo'lmasligi mumkin —
        # stilsiz ham jadval to'g'ri ishlaydi.
        pass
    header_cells = table.rows[0].cells
    for i, col in enumerate(IMPORT_HEADER):
        header_cells[i].text = col
    # Namuna qator
    sample_cells = table.add_row().cells
    for i, val in enumerate(IMPORT_SAMPLE_ROW):
        sample_cells[i].text = val

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = 'attachment; filename="olympy-savollar-namuna.docx"'
    return response


@api_view(['POST'])
@permission_classes([IsAuthenticated])
@parser_classes([MultiPartParser, FormParser])
@throttle_classes([AiQuestionRateThrottle])
def preview_pdf_questions(request):
    """POST /api/questions/pdf-preview/ — PDFdan savol ajratish task'ini boshlaydi.

    Avval bu view PDFni SINXRON tahlil qilardi (Gemini API 15-30 daqiqa) va
    Gunicorn worker'ni bloklardi. Endi Celery task'ni ishga tushiradi va darhol
    `task_id` qaytaradi; frontend `pdf-preview/<task_id>/status/` orqali polling
    qiladi. EAGER rejimda (Redis yo'q) task `delay()` ichida sinxron bajariladi.
    """
    import base64 as _b64
    import uuid

    from django.core.cache import cache

    from .tasks import process_pdf_questions_task

    center_id = request.data.get('center')
    if not center_id:
        return Response(
            {'detail': 'center majburiy'},
            status=http_status.HTTP_400_BAD_REQUEST,
        )
    if not _user_can_create_for_center(request.user, center_id):
        return Response(
            {'detail': "Savol yaratish uchun o'qituvchi/manager arizangiz tasdiqlanishi kerak"},
            status=http_status.HTTP_403_FORBIDDEN,
        )

    # Premium check
    from centers.models import EducationCenter
    center = EducationCenter.objects.filter(pk=center_id).first()
    if center and not center.is_premium:
        return Response(
            {
                'detail': "PDF tahlil orqali savollar ajratish faqat premium tashkilotlar uchun. Premium obunani faollashtiring.",
                'upgrade_required': True
            },
            status=http_status.HTTP_403_FORBIDDEN
        )

    pdf_file = request.FILES.get('pdf') or request.FILES.get('file') or request.FILES.get('document')
    if not pdf_file:
        return Response({'detail': 'PDF fayl yuboring'}, status=http_status.HTTP_400_BAD_REQUEST)
    filename = str(getattr(pdf_file, 'name', '') or '').lower()
    content_type = str(getattr(pdf_file, 'content_type', '') or '').lower()
    # Avval `AND` edi — content_type yoki kengaytma bittasi to'g'ri bo'lsa o'tib
    # ketardi. Endi `OR`: ikkalasi ham mos kelishi shart (PDF MIME + .pdf nomi).
    if content_type != 'application/pdf' or not filename.endswith('.pdf'):
        return Response({'detail': 'Faqat PDF fayl qabul qilinadi'}, status=http_status.HTTP_400_BAD_REQUEST)
    max_bytes = getattr(settings, 'AI_QUESTION_PDF_MAX_BYTES', 20 * 1024 * 1024)
    if pdf_file.size and pdf_file.size > max_bytes:
        return Response(
            {'detail': f"PDF juda katta. Limit: {max_bytes // (1024 * 1024)} MB"},
            status=http_status.HTTP_400_BAD_REQUEST,
        )
    pdf_bytes = pdf_file.read()
    if not pdf_bytes:
        return Response({'detail': "PDF fayl bo'sh"}, status=http_status.HTTP_400_BAD_REQUEST)
    # Magic bytes tekshiruvi — content_type/nom soxta bo'lishi mumkin, lekin
    # haqiqiy PDF har doim '%PDF' bilan boshlanadi.
    if pdf_bytes[:4] != b'%PDF':
        return Response(
            {'detail': "Fayl haqiqiy PDF emas (PDF imzosi topilmadi)"},
            status=http_status.HTTP_400_BAD_REQUEST,
        )

    task_id = str(uuid.uuid4())
    cache_key = f"pdf_questions:task:{task_id}"
    cache.set(cache_key, {'status': 'PENDING'}, timeout=900)
    # Celery argumenti JSON-serializable bo'lishi uchun PDF baytlarini base64 qilamiz.
    process_pdf_questions_task.delay(
        task_id,
        _b64.b64encode(pdf_bytes).decode('ascii'),
        request.data.get('subject') or '',
        request.data.get('difficulty') or 'medium',
        request.data.get('question_type') or 'multiple_choice',
    )
    return Response({'task_id': task_id}, status=http_status.HTTP_202_ACCEPTED)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
@parser_classes([MultiPartParser, FormParser])
@throttle_classes([AiQuestionRateThrottle])
def word_ai_preview(request):
    """POST /api/questions/word-ai-preview/ — Word matnidan AI savol ajratish task'ini boshlaydi.

    preview_pdf_questions bilan bir xil oqim: Celery task'ni ishga tushiradi va
    darhol `task_id` qaytaradi. Kesh kaliti PDF bilan bir xil pattern bo'lgani
    uchun (`pdf_questions:task:{task_id}`), frontend status'ni mavjud
    `pdf-preview/<task_id>/status/` endpointi orqali polling qiladi. Farqi:
    bu yerda .docx fayl qabul qilinadi va jadval emas, butun matn AI ga beriladi
    (import-word/ jadval formatini talab qiladi; bu esa erkin matnni ham ajratadi).
    EAGER rejimda (Redis yo'q) task `delay()` ichida sinxron bajariladi.
    """
    import base64 as _b64
    import uuid

    from django.core.cache import cache

    from .tasks import process_word_ai_questions_task

    center_id = request.data.get('center')
    if not center_id:
        return Response(
            {'detail': 'center majburiy'},
            status=http_status.HTTP_400_BAD_REQUEST,
        )
    if not _user_can_create_for_center(request.user, center_id):
        return Response(
            {'detail': "Savol yaratish uchun o'qituvchi/manager arizangiz tasdiqlanishi kerak"},
            status=http_status.HTTP_403_FORBIDDEN,
        )

    # Premium check — PDF tahlil bilan bir xil shart (AI tashqi xizmatga qimmat).
    from centers.models import EducationCenter
    center = EducationCenter.objects.filter(pk=center_id).first()
    if center and not center.is_premium:
        return Response(
            {
                'detail': "Word matnidan AI savollar ajratish faqat premium tashkilotlar uchun. Premium obunani faollashtiring.",
                'upgrade_required': True
            },
            status=http_status.HTTP_403_FORBIDDEN
        )

    word_file = request.FILES.get('word') or request.FILES.get('file') or request.FILES.get('document')
    if not word_file:
        return Response({'detail': "Fayl yuboring (.docx, .txt yoki .pdf)"}, status=http_status.HTTP_400_BAD_REQUEST)
    filename = str(getattr(word_file, 'name', '') or '').lower()
    # Qo'llab-quvvatlanadigan formatlar: .docx (python-docx), .txt (oddiy matn),
    # .pdf (pypdf/pdfplumber).
    allowed_ext = ('.docx', '.txt', '.pdf')
    if not filename.endswith(allowed_ext):
        return Response(
            {'detail': "Faqat .docx, .txt yoki .pdf fayl qabul qilinadi"},
            status=http_status.HTTP_400_BAD_REQUEST,
        )
    # Word import'i bilan bir xil limit (python-docx/pypdf butun hujjatni xotiraga yuklaydi).
    max_bytes = getattr(settings, 'AI_QUESTION_IMPORT_MAX_BYTES', 10 * 1024 * 1024)
    if word_file.size and word_file.size > max_bytes:
        return Response(
            {'detail': f"Fayl juda katta. Limit: {max_bytes // (1024 * 1024)} MB"},
            status=http_status.HTTP_400_BAD_REQUEST,
        )
    word_bytes = word_file.read()
    if not word_bytes:
        return Response({'detail': "Fayl bo'sh"}, status=http_status.HTTP_400_BAD_REQUEST)
    # Magic bytes bilan format imzosini tekshiramiz (kengaytma yolg'on bo'lishi mumkin).
    if filename.endswith('.docx') and word_bytes[:2] != b'PK':
        # .docx — ZIP arxivi, har doim 'PK' (PK\x03\x04) bilan boshlanadi.
        return Response(
            {'detail': "Fayl haqiqiy .docx emas (Word imzosi topilmadi)"},
            status=http_status.HTTP_400_BAD_REQUEST,
        )
    if filename.endswith('.pdf') and word_bytes[:4] != b'%PDF':
        # .pdf — har doim '%PDF' imzosi bilan boshlanadi.
        return Response(
            {'detail': "Fayl haqiqiy .pdf emas (PDF imzosi topilmadi)"},
            status=http_status.HTTP_400_BAD_REQUEST,
        )

    # Fayl turini task'ga uzatamiz — qaysi parser ishlatilishini aniqlash uchun.
    if filename.endswith('.txt'):
        file_kind = 'txt'
    elif filename.endswith('.pdf'):
        file_kind = 'pdf'
    else:
        file_kind = 'docx'

    task_id = str(uuid.uuid4())
    cache_key = f"pdf_questions:task:{task_id}"
    cache.set(cache_key, {'status': 'PENDING'}, timeout=900)
    # Celery argumenti JSON-serializable bo'lishi uchun fayl baytlarini base64 qilamiz.
    process_word_ai_questions_task.delay(
        task_id,
        _b64.b64encode(word_bytes).decode('ascii'),
        request.data.get('subject') or '',
        request.data.get('difficulty') or 'medium',
        request.data.get('question_type') or 'multiple_choice',
        file_kind,
    )
    return Response({'task_id': task_id}, status=http_status.HTTP_202_ACCEPTED)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def pdf_preview_status(request, task_id):
    """GET /api/questions/pdf-preview/<task_id>/status/ — PDF task holatini qaytaradi.

    Javob shakllari:
      {'status': 'PENDING'}
      {'status': 'COMPLETED', 'questions': [...], 'provider': ..., ...}
      {'status': 'FAILED', 'detail': ...}
    """
    from django.core.cache import cache

    state = cache.get(f"pdf_questions:task:{task_id}")
    if not state:
        return Response(
            {'status': 'FAILED', 'detail': "Vazifa topilmadi yoki muddati o'tgan"},
            status=http_status.HTTP_404_NOT_FOUND,
        )
    status_value = state.get('status')
    if status_value == 'COMPLETED':
        result = state.get('result') or {}
        return Response({'status': 'COMPLETED', **result})
    if status_value == 'FAILED':
        # Frontend xato xabarini `detail`/`error` orqali ko'rsatadi.
        return Response({
            'status': 'FAILED',
            'detail': state.get('error') or "PDFdan savollarni ajratib bo'lmadi",
            'pdf_text_chars': state.get('pdf_text_chars', 0),
            'page_count': state.get('page_count', 0),
            'used_pdf_vision': bool(state.get('used_pdf_vision')),
        })
    return Response({'status': 'PENDING'})


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def question_analytics(request):
    """GET /api/questions/analytics/?center=<id>
    Markazdagi savollar bo'yicha noto'g'ri javob statistikasi.

    Har bir savol uchun: question_id, text (qisqartirilgan, 80 belgi),
    subject, total_attempts, wrong_count, wrong_rate (%).
    Faqat wrong_rate >= 30% va total_attempts >= 3 bo'lganlar.
    wrong_rate kamayish tartibida, max 50 ta.
    """
    raw_center = request.query_params.get('center')
    if not raw_center:
        return Response(
            {'detail': 'center query parametri majburiy'},
            status=http_status.HTTP_400_BAD_REQUEST,
        )
    try:
        center_id = int(raw_center)
    except (TypeError, ValueError):
        return Response(
            {'detail': "center parametri son bo'lishi kerak"},
            status=http_status.HTTP_400_BAD_REQUEST,
        )
    if not _user_can_create_for_center(request.user, center_id):
        return Response(
            {'detail': 'Forbidden'},
            status=http_status.HTTP_403_FORBIDDEN,
        )

    from attempts.models import TestAttempt

    questions = list(
        Question.objects.filter(center_id=center_id)
        .only('id', 'text', 'subject', 'correct_answer')
    )
    if not questions:
        return Response([])

    qmap = {q.id: q for q in questions}
    total_count = {q.id: 0 for q in questions}
    wrong_count = {q.id: 0 for q in questions}

    answers_iter = (
        TestAttempt.objects
        .filter(olympiad__center_id=center_id, disqualified=False)
        .values_list('answers', flat=True)
    )

    for ans in answers_iter:
        if not isinstance(ans, dict):
            continue
        for k, v in ans.items():
            try:
                qid = int(k)
            except (TypeError, ValueError):
                continue
            q = qmap.get(qid)
            if not q:
                continue
            total_count[qid] = total_count.get(qid, 0) + 1
            try:
                chosen = int(v)
            except (TypeError, ValueError):
                # Javob berilmagan/noto'g'ri formatda — noto'g'ri deb hisoblaymiz.
                wrong_count[qid] = wrong_count.get(qid, 0) + 1
                continue
            if chosen != q.correct_answer:
                wrong_count[qid] = wrong_count.get(qid, 0) + 1

    rows = []
    for q in questions:
        total = total_count.get(q.id, 0)
        wrong = wrong_count.get(q.id, 0)
        if total < 3:
            continue
        rate = round((wrong / total) * 100.0, 1) if total else 0.0
        if rate < 30.0:
            continue
        text_short = (q.text or '')[:80]
        if q.text and len(q.text) > 80:
            text_short = text_short.rstrip() + '…'
        rows.append({
            'question_id': q.id,
            'text': text_short,
            'subject': q.subject or '',
            'total_attempts': total,
            'wrong_count': wrong,
            'wrong_rate': rate,
        })

    rows.sort(key=lambda r: -r['wrong_rate'])
    return Response(rows[:50])


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def olympiad_questions(request, olympiad_id):
    from datetime import timedelta

    from django.utils import timezone

    from olympiads.models import Olympiad
    from olympiads.services import (
        maybe_finish_expired_olympiad,
        user_can_participate_in_event,
    )

    olympiad = get_object_or_404(Olympiad, pk=olympiad_id)
    # Celery worker bo'lmagan muhitda muddati o'tgan olimpiadani lazy yopish.
    # Status o'zgargan bo'lsa pastdagi tekshiruvlar to'g'ri ishlaydi.
    try:
        maybe_finish_expired_olympiad(olympiad)
        olympiad.refresh_from_db()
    except Exception:
        pass
    if olympiad.status != Olympiad.STATUS_ACTIVE:
        # Status'ga qarab aniqroq xabar — student "Olimpiada faol emas"
        # ko'rganida tushunmasdi (yakunlanganmi yoki hali boshlanmaganmi?).
        if olympiad.status == Olympiad.STATUS_FINISHED:
            detail = "Olimpiada yakunlangan"
        elif olympiad.status == Olympiad.STATUS_DRAFT:
            detail = "Olimpiada hali nashr qilinmagan"
        else:
            detail = "Olimpiada faol emas"
        return Response({'detail': detail}, status=http_status.HTTP_403_FORBIDDEN)
    # Time-window check (timezone-aware): the celery finisher may not have run
    # yet, so don't trust status alone.
    now = timezone.now()
    if olympiad.start_datetime and now < olympiad.start_datetime:
        return Response(
            {'detail': "Olimpiada hali boshlanmagan"},
            status=http_status.HTTP_400_BAD_REQUEST,
        )
    if olympiad.start_datetime and olympiad.duration_minutes:
        end_time = olympiad.start_datetime + timedelta(minutes=olympiad.duration_minutes)
        if now > end_time:
            return Response(
                {'detail': "Olimpiada vaqti tugagan"},
                status=http_status.HTTP_400_BAD_REQUEST,
            )
    if not user_can_participate_in_event(request.user, olympiad):
        return Response({'detail': 'Forbidden'}, status=http_status.HTTP_403_FORBIDDEN)
    from attempts.models import TestAttempt
    from attempts.session_utils import (
        get_or_create_test_session,
        questions_payload,
        session_is_expired,
        session_timing_payload,
    )

    if TestAttempt.objects.filter(user=request.user, olympiad=olympiad).exists():
        return Response(
            {'detail': "Siz bu olimpiadaga allaqachon qatnashgansiz"},
            status=http_status.HTTP_400_BAD_REQUEST,
        )
    session = get_or_create_test_session(request.user, olympiad)
    if session.status == getattr(session, 'STATUS_DISQUALIFIED', 'disqualified'):
        return Response(
            {'detail': "Siz cheating qildingiz. Olimpiada yakunlandi."},
            status=http_status.HTTP_403_FORBIDDEN,
        )
    if session_is_expired(session, olympiad):
        return Response(
            {'detail': "Test vaqti tugagan"},
            status=http_status.HTTP_400_BAD_REQUEST,
        )
    # Avval bu yerda faqat questions arrayi qaytarilardi va frontend lokal
    # DURATION dan teskari sanardi. Endi questions ham, server timing'i ham
    # qaytariladi — bu frontend va server vaqti orasidagi drift'ni yo'qotadi.
    # Backward-compat: response.questions arrayi avvalgi roli bilan bir xil.
    all_questions = questions_payload(session, olympiad)

    # Cheating-himoya: savollarni bitta-bitta yuklash. `q` param berilsa,
    # faqat o'sha indeksdagi savol qaytariladi — shu tariqa bir vaqtda
    # barcha savollarni ko'chirib olish (scrape) qiyinlashadi.
    q_param = request.query_params.get('q')
    if q_param is not None:
        try:
            q_index = int(q_param)
        except (TypeError, ValueError):
            return Response(
                {'detail': "Noto'g'ri savol indeksi"},
                status=http_status.HTTP_400_BAD_REQUEST,
            )
        if q_index < 0 or q_index >= len(all_questions):
            return Response(
                {'detail': "Savol indeksi diapazondan tashqarida"},
                status=http_status.HTTP_400_BAD_REQUEST,
            )
        return Response({
            'questions': [all_questions[q_index]],
            'question_index': q_index,
            'total_questions': len(all_questions),
            'session': session_timing_payload(session, olympiad),
        })

    return Response({
        'questions': all_questions,
        'session': session_timing_payload(session, olympiad),
    })


@api_view(['POST'])
@permission_classes([IsAuthenticated])
@throttle_classes([CodeReviewRateThrottle])
def code_review(request):
    """POST /api/questions/code-review/ — IT kod savolini AI bilan baholaydi.

    Body: { question_id, submitted_code, language }. O'quvchi test paytida
    kodini sinash uchun AI feedback oladi (to'g'rilik 0-100, xatolar, tavsiya).
    Bu yerda natija SAQLANMAYDI — yakuniy kod javob va AI bahosi submit
    paytida CodeSubmission'ga yoziladi. Rate limit: 10/hour (code_review).
    """
    question_id = request.data.get('question_id')
    submitted_code = request.data.get('submitted_code') or ''
    language = (request.data.get('language') or '').strip().lower()
    if not question_id:
        return Response({'detail': 'question_id majburiy'}, status=http_status.HTTP_400_BAD_REQUEST)
    if not str(submitted_code).strip():
        return Response({'detail': 'Kod bo\'sh bo\'lishi mumkin emas'}, status=http_status.HTTP_400_BAD_REQUEST)

    question = get_object_or_404(Question, pk=question_id)
    if question.question_type != Question.QUESTION_TYPE_CODE:
        return Response(
            {'detail': 'Bu kod (IT) savoli emas'},
            status=http_status.HTTP_400_BAD_REQUEST,
        )

    result = review_code_submission(
        question_text=question.text,
        submitted_code=submitted_code,
        language=language or question.programming_language,
        expected_output=question.expected_output or '',
    )
    return Response({
        'score': result.get('score'),
        'review': result.get('review') or '',
    })


def _normalize_output(text):
    """Test case taqqoslash uchun stdout/expected normalizatsiyasi.

    Trailing whitespace va satr oxiridagi bo'shliqlar Judge0 chiqishida tez-tez
    farq qiladi (masalan oxirgi '\\n'). Har bir satrning o'ng tomonini va butun
    matnning oxirini tozalaymiz — bu 'to'g'ri javob bo'sh joy tufayli xato'
    holatini oldini oladi.
    """
    if text is None:
        return ''
    lines = str(text).replace('\r\n', '\n').replace('\r', '\n').split('\n')
    return '\n'.join(line.rstrip() for line in lines).rstrip()


@api_view(['POST'])
@permission_classes([IsAuthenticated])
@throttle_classes([AiExplainRateThrottle])
def explain_question(request, question_id):
    """POST /api/questions/<id>/explain/

    Savol uchun yechim tushuntirishini qaytaradi. Agar tushuntirish bazada
    saqlanmagan bo'lsa, Gemini AI yordamida generatsiya qilinadi va keshlanadi.

    Rate limit: 'ai' scope (settings.REST_FRAMEWORK.DEFAULT_THROTTLE_RATES) —
    tashqi Gemini API'ga qimmat va sekin murojaat qiladi, abuse'dan himoyalanadi.
    """
    question = get_object_or_404(Question, pk=question_id)
    # Center ownership tekshiruvi — istalgan foydalanuvchi boshqa markazga
    # tegishli savol ID'sini yuborib AI tushuntirishni olmasligi uchun. Faqat
    # shu markazda tasdiqlangan o'qituvchi/menejer/egasi (yoki platforma admini)
    # tushuntirishni ola oladi.
    if not _user_can_create_for_center(request.user, question.center_id):
        return Response({'detail': 'Forbidden'}, status=http_status.HTTP_403_FORBIDDEN)
    if question.explanation and question.explanation.strip():
        return Response({'explanation': question.explanation})

    explanation_text = explain_question_ai(
        question_text=question.text,
        options=question.options or [],
        correct_idx=question.correct_answer,
        subject=question.subject or '',
    )

    if explanation_text and "generatsiya qilinmadi" not in explanation_text:
        question.explanation = explanation_text
        question.save(update_fields=['explanation'])

    return Response({'explanation': explanation_text})


@api_view(['POST'])
@permission_classes([IsAuthenticated])
@throttle_classes([CodeRunRateThrottle])
def run_code_start_view(request):
    """POST /api/questions/run-code/start/ — Asynchronously triggers a Judge0 run task."""
    import uuid
    from django.core.cache import cache
    from .judge0_service import is_supported
    from .tasks import run_code_async_task

    source_code = request.data.get('source_code') or ''
    language = (request.data.get('language') or '').strip().lower()
    stdin = request.data.get('stdin') or ''
    question_id = request.data.get('question_id')

    if not str(source_code).strip():
        return Response(
            {'detail': "Kod bo'sh bo'lishi mumkin emas"},
            status=http_status.HTTP_400_BAD_REQUEST,
        )
    if not language:
        return Response(
            {'detail': "Dasturlash tili majburiy"},
            status=http_status.HTTP_400_BAD_REQUEST,
        )
    if not is_supported(language):
        return Response(
            {'detail': f"'{language}' tili qo'llab-quvvatlanmaydi"},
            status=http_status.HTTP_400_BAD_REQUEST,
        )

    task_id = str(uuid.uuid4())
    cache.set(f"run_code:task:{task_id}", {'status': 'PENDING'}, timeout=300)
    
    run_code_async_task.delay(task_id, source_code, language, stdin, question_id)
    
    return Response({'task_id': task_id}, status=http_status.HTTP_202_ACCEPTED)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def run_code_status_view(request, task_id):
    """GET /api/questions/run-code/status/<task_id>/ — Returns the status of the run code task."""
    from django.core.cache import cache
    state = cache.get(f"run_code:task:{task_id}")
    if not state:
        return Response(
            {'status': 'FAILED', 'error': "Vazifa topilmadi yoki muddati o'tgan"},
            status=http_status.HTTP_404_NOT_FOUND,
        )
    return Response(state)
